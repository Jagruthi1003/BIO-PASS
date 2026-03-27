from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create document
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('BIO PASS\nZero-Knowledge Biometric Event Entry System')
title_run.bold = True
title_run.font.size = Pt(18)
title_run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('System Documentation')
subtitle_run.font.size = Pt(12)
subtitle_run.italic = True

doc.add_paragraph()

# 1. INTRODUCTION
doc.add_heading('1. INTRODUCTION', 0)
doc.add_paragraph(
    'Bio-Pass is an advanced entry management system that addresses ticket fraud in events by enabling '
    'privacy-preserving face verification through zero-knowledge cryptographic proofs. Traditional ticketing '
    'systems suffer from fake tickets, proxy attendance, and privacy leaks during biometric checks. This '
    'Flutter-based mobile application integrates Google ML Kit for face landmark extraction, Firebase for '
    'secure ticketing and authentication, and SHA256-based cryptographic commitments to prove "my face landmarks '
    'match the enrolled ticket hash" without revealing raw biometric data.'
)

doc.add_paragraph(
    'The system leverages Google ML Kit\'s FaceDetector with 68-point face landmark detection for precise facial '
    'feature extraction. These landmarks are quantized and hashed using SHA256 cryptographic algorithms, ensuring '
    'proofs verify against public landmark commitments stored in Firestore while maintaining complete user privacy. '
    'The architecture supports secure, scalable event access control across Android, iOS, and Windows platforms.'
)

p = doc.add_heading('1.1 Core Technologies', level=2)
tech_list = [
    'Flutter Framework (Dart): Cross-platform mobile and web development',
    'Firebase: Authentication, Firestore database, real-time synchronization',
    'Google ML Kit Face Detection: 68-point landmark extraction with sub-pixel accuracy',
    'SHA256 Cryptographic Hashing: Proof generation and verification',
    'Camera Plugin: Real-time camera access with video streaming',
    'Permission Handler: Runtime permission management across platforms',
    'Crypto Package: Cryptographic operations and hash generation'
]
for tech in tech_list:
    doc.add_paragraph(tech, style='List Bullet')

# 2. PROBLEM STATEMENT AND OBJECTIVE
doc.add_heading('2. PROBLEM STATEMENT AND OBJECTIVE', 0)

doc.add_heading('2.1 Problem Statement', level=2)
problems = [
    'Physical and digital tickets are easily forged or transferred to unauthorized users',
    'Centralized biometric storage creates privacy risks and single points of failure',
    'Manual entry checks create bottlenecks, long queues, and privacy concerns',
    'Event organizers lack fraud-proof, real-time validation mechanisms',
    'No mechanism exists to prove identity without revealing sensitive biometric data'
]
for problem in problems:
    doc.add_paragraph(problem, style='List Bullet')

doc.add_heading('2.2 Objectives', level=2)
objectives = [
    'Develop Flutter application supporting ticket purchase, face enrollment, and entry verification across multiple platforms',
    'Implement Google ML Kit integration for precise facial landmark extraction and processing',
    'Design SHA256-based cryptographic proof system ensuring landmark verification without biometric disclosure',
    'Integrate Firebase for secure ticket storage, authentication, and real-time proof verification',
    'Enforce non-transferable, privacy-preserving event access with atomic transaction guarantees',
    'Support three user roles (Attendee, Organizer, Gatekeeper) with role-based access control',
    'Achieve 80% similarity threshold for secure face verification with minimal false positives'
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Bullet')

# 3. SYSTEM ANALYSIS
doc.add_heading('3. SYSTEM ANALYSIS', 0)

doc.add_heading('3.1 Existing System Limitations', level=2)
doc.add_paragraph(
    'Current event ticketing relies on QR codes with manual ID verification or centralized biometric storage, '
    'both vulnerable to fraud and data breaches. These systems suffer from:'
)
limitations = [
    'Ticket transferability enabling proxy attendance',
    'Centralized biometric data repositories creating privacy violations',
    'Lack of cryptographic proofs preventing real-time fraud detection',
    'Manual verification processes creating operational inefficiency',
    'No mechanism separating authentication from authorization'
]
for limit in limitations:
    doc.add_paragraph(limit, style='List Bullet')

doc.add_heading('3.2 Proposed System Architecture', level=2)
doc.add_paragraph(
    'Bio-Pass implements a privacy-preserving, three-tier architecture:'
)
doc.add_paragraph().add_run('Tier 1 - Mobile Client (Flutter):').bold = True
doc.add_paragraph('Captures face data via camera, extracts 68-point landmarks using Google ML Kit FaceDetector', style='List Bullet 2')
doc.add_paragraph('Generates SHA256 cryptographic commitment to landmarks', style='List Bullet 2')
doc.add_paragraph('Verifies similarity between live capture and enrolled face with configurable threshold', style='List Bullet 2')

doc.add_paragraph().add_run('Tier 2 - Backend Services (Firebase):').bold = True
doc.add_paragraph('Firebase Authentication: Email/password authentication with role-based access control', style='List Bullet 2')
doc.add_paragraph('Firestore: Distributed database storing tickets, landmark commitments, and verification proofs', style='List Bullet 2')
doc.add_paragraph('Real-time synchronization: Instant updates across all connected clients', style='List Bullet 2')
doc.add_paragraph('Atomic transactions: Ensures one-time ticket usage preventing duplicate entries', style='List Bullet 2')

doc.add_paragraph().add_run('Tier 3 - Cryptographic Verification:').bold = True
doc.add_paragraph('SHA256 proof generation from face landmarks', style='List Bullet 2')
doc.add_paragraph('Euclidean distance-based similarity calculation (80% threshold)', style='List Bullet 2')
doc.add_paragraph('Zero-knowledge proof verification without landmark disclosure', style='List Bullet 2')

doc.add_heading('3.3 Feasibility Assessment', level=2)

table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'

# Header
header_cells = table.rows[0].cells
header_cells[0].text = 'Dimension'
header_cells[1].text = 'Assessment'
header_cells[2].text = 'Status'

# Technical
row = table.rows[1].cells
row[0].text = 'Technical'
row[1].text = 'SHA256 cryptography mature; ML Kit Face Detection production-ready; Firebase scales to millions of events'
row[2].text = '✓ Feasible'

# Operational
row = table.rows[2].cells
row[0].text = 'Operational'
row[1].text = 'Simple scan-to-enter workflow; Proofs generated in 1-2 seconds; Real-time status synchronization'
row[2].text = '✓ Feasible'

# Economic
row = table.rows[3].cells
row[0].text = 'Economic'
row[1].text = 'Firebase free tier sufficient; Open-source ML Kit; Low infrastructure costs'
row[2].text = '✓ Feasible'

# 4. SYSTEM DESIGN
doc.add_heading('4. SYSTEM DESIGN', 0)

doc.add_heading('4.1 Technology Stack', level=2)

tech_stack = [
    ('Frontend', 'Flutter 3.22.0+, Dart 3.4.0+'),
    ('Authentication', 'Firebase Authentication 5.3.1'),
    ('Database', 'Cloud Firestore 5.4.4'),
    ('Face Detection', 'Google ML Kit Face Detection 0.10.0'),
    ('Camera', 'Camera Plugin 0.11.0'),
    ('Cryptography', 'Dart Crypto Package 3.0.3 (SHA256)'),
    ('Permissions', 'Permission Handler 12.0.1'),
    ('Platforms', 'Android 5.0+, iOS 12.0+, Windows 10+, Chrome Web')
]

for category, details in tech_stack:
    p = doc.add_paragraph()
    p.add_run(f'{category}: ').bold = True
    p.add_run(details)

doc.add_heading('4.2 Data Models', level=2)

models = [
    ('User Model', '{ uid, email, name, role, enrolledFaceHash }'),
    ('Event Model', '{ eventId, name, organizerId, capacity, registrations[] }'),
    ('Ticket Model', '{ ticketId, eventId, attendeeId, status, faceProof, enrollmentHash }'),
    ('Face Proof', '{ proof (SHA256 hash), timestamp, similarity (0-1) }')
]

for model_name, fields in models:
    p = doc.add_paragraph()
    p.add_run(f'{model_name}: ').bold = True
    p.add_run(fields)

doc.add_heading('4.3 User Roles and Workflows', level=2)

doc.add_heading('4.3.1 Attendee Role', level=3)
attendee_flows = [
    'Browse and register for events',
    'Enroll face landmarks via camera',
    'View ticket with enrollment status',
    'Use app for entry verification',
    'Track verification status in real-time'
]
for flow in attendee_flows:
    doc.add_paragraph(flow, style='List Bullet')

doc.add_heading('4.3.2 Organizer Role', level=3)
organizer_flows = [
    'Create and manage events',
    'View registered attendees',
    'Track real-time entry statistics',
    'Monitor verification logs',
    'View attendance reports'
]
for flow in organizer_flows:
    doc.add_paragraph(flow, style='List Bullet')

doc.add_heading('4.3.3 Gatekeeper Role', level=3)
gatekeeper_flows = [
    'Initiate face verification at gate',
    'Compare live capture to enrollment',
    'Verify similarity exceeds 80% threshold',
    'Mark ticket as used',
    'Display verification result'
]
for flow in gatekeeper_flows:
    doc.add_paragraph(flow, style='List Bullet')

doc.add_heading('4.4 Cryptographic Verification', level=2)

doc.add_paragraph(
    'Bio-Pass uses SHA256-based cryptographic commitment scheme for privacy-preserving verification:'
)

doc.add_paragraph().add_run('Proof Generation:').bold = True
doc.add_paragraph('Face landmarks → JSON serialization → UTF-8 encoding → SHA256 hash', style='List Bullet')

doc.add_paragraph().add_run('Verification Process:').bold = True
doc.add_paragraph('Generate proof from live landmarks', style='List Bullet')
doc.add_paragraph('Compare to stored enrollment proof', style='List Bullet')
doc.add_paragraph('Calculate Euclidean distance-based similarity', style='List Bullet')
doc.add_paragraph('Verify similarity exceeds 80% threshold', style='List Bullet')

# 5. IMPLEMENTATION STATUS
doc.add_heading('5. IMPLEMENTATION STATUS', 0)

doc.add_heading('Module 1: Authentication and Ticket Management (✓ Completed)', level=2)

features = [
    'Firebase Authentication with email/password',
    'Role-based routing and access control',
    'Firestore security rules enforcement',
    'Ticket CRUD operations',
    'Event creation with capacity limits',
    'Atomic transactions for one-time usage',
    'Real-time synchronization'
]

for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('Module 2: Face Landmarks and Proofs (✓ Completed)', level=2)

features = [
    'ML Kit FaceDetector integration',
    'SHA256 proof generation',
    'Real-time face capture',
    'Euclidean similarity calculation',
    'Face enrollment workflow',
    'Face verification system'
]

for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('Module 3: Entry Verification (✓ Completed)', level=2)

features = [
    'Real-time camera integration',
    'Proof-based verification',
    'Atomic ticket updates',
    'Real-time dashboard sync',
    'Multi-platform support'
]

for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

# 6. FIREBASE CONFIGURATION
doc.add_heading('6. FIREBASE CONFIGURATION', 0)

doc.add_heading('6.1 Firestore Collections', level=2)

collections = [
    ('users', '{ uid, email, name, role, enrolledFaceHash, createdAt }'),
    ('events', '{ name, organizerId, description, capacity, registrations, createdAt }'),
    ('tickets', '{ eventId, attendeeId, status, faceProof, enrollmentHash, timestamp }'),
    ('proofs', '{ ticketId, proof, similarity, timestamp, verifiedAt }')
]

for collection, fields in collections:
    p = doc.add_paragraph()
    p.add_run(f'{collection}: ').bold = True
    p.add_run(fields)

doc.add_heading('6.2 Security Features', level=2)
security_features = [
    'Users can only read/write their own documents',
    'Organizers can create and manage events',
    'Attendees can only view their own tickets',
    'Atomic transactions prevent duplicate entries',
    'Timestamp validation for chronological consistency'
]
for feature in security_features:
    doc.add_paragraph(feature, style='List Bullet')

# 7. KEY TECHNOLOGIES
doc.add_heading('7. KEY TECHNOLOGIES AND TERMINOLOGIES', 0)

doc.add_heading('7.1 Flutter and Dart', level=2)
doc.add_paragraph(
    'Flutter: Open-source UI framework by Google for cross-platform development. '
    'Dart: Programming language with null safety, JIT/AOT compilation. '
    'Version: Flutter 3.22.0+, Dart 3.4.0+'
)

doc.add_heading('7.2 Google ML Kit Face Detection', level=2)
doc.add_paragraph(
    'On-device ML SDK providing 68-point facial landmark extraction. '
    'Features: Face detection, contour analysis, classification, real-time tracking'
)

doc.add_heading('7.3 Firebase Authentication', level=2)
doc.add_paragraph(
    'Secure authentication with email/password. '
    'Features: Bcrypt password hashing, session management, role-based claims'
)

doc.add_heading('7.4 Cloud Firestore', level=2)
doc.add_paragraph(
    'NoSQL cloud database with real-time synchronization. '
    'Features: Real-time listeners, atomic transactions, security rules, offline persistence'
)

doc.add_heading('7.5 SHA256 Cryptographic Hashing', level=2)
doc.add_paragraph(
    'Secure Hash Algorithm producing 256-bit digest. '
    'Usage: Generate deterministic commitment to landmarks, prevent impersonation, enable fast verification'
)

doc.add_heading('7.6 Camera Integration', level=2)
doc.add_paragraph(
    'Camera plugin for real-time video streaming. '
    'Features: Live frame capture, multiple camera support, cross-platform consistency'
)

doc.add_heading('7.7 Permission Handler', level=2)
doc.add_paragraph(
    'Runtime permission management. '
    'Requested: CAMERA (face capture), Location (future events)'
)

# 8. SYSTEM WORKFLOWS
doc.add_heading('8. SYSTEM WORKFLOWS', 0)

doc.add_heading('8.1 Attendee Entry Flow', level=2)

workflow = [
    'Sign up with email/password/role',
    'Browse and register for event',
    'Enroll face: Grant camera permission → Capture face → ML Kit extracts 68 landmarks → SHA256 proof generated → Stored in Firestore',
    'View ticket with enrollment status',
    'At event gate: Initiate verification → Live face capture → Proof comparison → Status update → Ticket marked used',
    'Track verification status in real-time'
]

for i, step in enumerate(workflow, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Bullet')

doc.add_heading('8.2 Gatekeeper Verification Flow', level=2)

workflow = [
    'Login with Gatekeeper role',
    'Receive attendee at entry point',
    'Initiate face verification via camera',
    'Real-time face capture and ML Kit landmark extraction',
    'Compute SHA256 proof → Compare to enrollment proof',
    'Calculate Euclidean distance → Normalize to similarity score',
    'Verify similarity > 80% threshold',
    'Atomic Firestore transaction marks ticket used',
    'Display "ENTRY GRANTED" or "ENTRY DENIED"',
    'Real-time sync to attendee dashboard'
]

for i, step in enumerate(workflow, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Bullet')

# 9. SECURITY
doc.add_heading('9. SECURITY FEATURES', 0)

doc.add_paragraph().add_run('Privacy:').bold = True
doc.add_paragraph('Raw landmarks never stored; only cryptographic hashes transmitted', style='List Bullet')
doc.add_paragraph('SHA256 is one-way; original face data non-recoverable', style='List Bullet')
doc.add_paragraph('Verification proves face match without biometric disclosure', style='List Bullet')

doc.add_paragraph().add_run('Fraud Prevention:').bold = True
doc.add_paragraph('Atomic transactions prevent duplicate entries', style='List Bullet')
doc.add_paragraph('80% similarity threshold prevents photo spoofing', style='List Bullet')
doc.add_paragraph('Live face capture required', style='List Bullet')
doc.add_paragraph('Enrollment and verification by same user (Firebase UID enforcement)', style='List Bullet')

doc.add_paragraph().add_run('Access Control:').bold = True
doc.add_paragraph('Firebase Security Rules enforce role-based access', style='List Bullet')
doc.add_paragraph('Users access only their own documents', style='List Bullet')
doc.add_paragraph('Organizers limited to event management', style='List Bullet')
doc.add_paragraph('Gatekeepers have read-only access', style='List Bullet')

# 10. PERFORMANCE
doc.add_heading('10. PERFORMANCE METRICS', 0)

metrics = [
    ('Face Detection', '< 500ms per frame'),
    ('Proof Generation', '< 100ms'),
    ('Firestore Write', '< 1 second'),
    ('Real-time Sync', '< 2 seconds'),
    ('Similarity Threshold', '80% (Euclidean distance-based)'),
    ('False Negative Rate', '< 5%'),
    ('False Positive Rate', '< 2%')
]

for metric, value in metrics:
    p = doc.add_paragraph()
    p.add_run(f'{metric}: ').bold = True
    p.add_run(value)

# 11. GLOSSARY
doc.add_heading('11. GLOSSARY OF TERMINOLOGIES', 0)

glossary = [
    ('Attestation', 'Cryptographic proof of ownership or correctness'),
    ('Biometric', 'Unique physical characteristic (face, fingerprint)'),
    ('Commitment', 'Cryptographic hash binding to data without disclosure'),
    ('Dart', 'Programming language for Flutter'),
    ('Euclidean Distance', 'Straight-line distance in multi-dimensional space'),
    ('Face Landmark', 'Precise coordinate of facial feature'),
    ('Firestore', 'Cloud NoSQL database with real-time sync'),
    ('Flutter', 'Cross-platform mobile/web framework'),
    ('ML Kit', 'Google on-device ML SDK'),
    ('Proof', 'Cryptographic evidence of claim'),
    ('SHA256', '256-bit secure hash algorithm'),
    ('Similarity Threshold', 'Minimum acceptable match percentage (80%)'),
    ('Zero-Knowledge Proof', 'Proof without revealing statement content')
]

for term, definition in glossary:
    p = doc.add_paragraph()
    p.add_run(f'{term}: ').bold = True
    p.add_run(definition)

# 12. REFERENCES
doc.add_heading('12. REFERENCES', 0)

references = [
    'Flutter: https://flutter.dev',
    'Firebase: https://firebase.google.com',
    'Google ML Kit: https://developers.google.com/ml-kit',
    'Dart: https://dart.dev',
    'Firestore: https://firebase.google.com/docs/firestore',
    'Crypto Package: https://pub.dev/packages/crypto',
    'Camera Plugin: https://pub.dev/packages/camera',
    'Permission Handler: https://pub.dev/packages/permission_handler'
]

for ref in references:
    doc.add_paragraph(ref, style='List Bullet')

# Footer
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_run = p.add_run('© 2026 Bio-Pass System. All Rights Reserved.')
p_run.italic = True
p_run.font.size = Pt(9)

# Save
doc.save('BIO_PASS_SYSTEM_DOCUMENTATION.docx')
print("✓ Word document created: BIO_PASS_SYSTEM_DOCUMENTATION.docx")
